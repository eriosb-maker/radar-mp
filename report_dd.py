"""
Generador de informes DOCX de Due Diligence de Contratistas.
Estilo: firma de abogados de alto estándar (Kirkland & Ellis / Carey).
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─── Paleta ────────────────────────────────────────────────────────
NAVY      = RGBColor(0x19, 0x2A, 0x56)
CHAMPAGNE = RGBColor(0xC8, 0xA9, 0x51)
DARK_GRAY = RGBColor(0x2C, 0x2C, 0x2C)
MID_GRAY  = RGBColor(0x60, 0x60, 0x60)
RED_FLAG  = RGBColor(0xC0, 0x39, 0x2B)
ORANGE    = RGBColor(0xE6, 0x7E, 0x22)
GREEN     = RGBColor(0x27, 0xAE, 0x60)


def _set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _nivel_color(nivel: str) -> RGBColor:
    return {"ALTO": RED_FLAG, "MEDIO": ORANGE, "BAJO": GREEN}.get(nivel, MID_GRAY)


def fmt_clp(n) -> str:
    if n is None:
        return "No informado"
    return f"$ {int(n):,}".replace(",", ".") + " CLP"


def fmt_fecha(iso: str | None) -> str:
    if not iso:
        return "—"
    return datetime.fromisoformat(iso).strftime("%d/%m/%Y")


# ─── Estilos de párrafo ─────────────────────────────────────────────

def _h1(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto.upper())
    run.bold       = True
    run.font.size  = Pt(13)
    run.font.color.rgb = NAVY
    # Línea inferior
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "C8A951")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _h2(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(texto)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_GRAY


def _body(doc, texto, italic=False):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        if italic:
            run.italic = True
    return p


def _kv(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = MID_GRAY
    r2 = p.add_run(str(value))
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY


# ─── Documento principal ─────────────────────────────────────────────

def generar_docx(informe: dict) -> bytes:
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2.5))

    prov    = informe["proveedor"]
    metr    = informe["metricas"]
    flags   = informe["red_flags"]
    analisis = informe["analisis_ia"]

    # ── PORTADA ──────────────────────────────────────────────────────
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run("INFORME DE DUE DILIGENCE")
    r.bold           = True
    r.font.size      = Pt(20)
    r.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Contratista en Mercado Público")
    r2.font.size      = Pt(12)
    r2.font.color.rgb = CHAMPAGNE

    doc.add_paragraph()

    # Tabla de identificación
    tabla_id = doc.add_table(rows=5, cols=2)
    tabla_id.style = "Table Grid"
    datos_id = [
        ("Proveedor",           prov["nombre"] or "—"),
        ("RUT",                 prov["rut"]),
        ("Código Mercado Público", prov["codigo"] or "—"),
        ("Fecha del informe",   datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Período analizado",   f"{fmt_fecha(metr['primer_contrato'])} — {fmt_fecha(metr['ultimo_contrato'])}"),
    ]
    for i, (lbl, val) in enumerate(datos_id):
        tabla_id.cell(i, 0).text = lbl
        tabla_id.cell(i, 1).text = val
        _set_cell_bg(tabla_id.cell(i, 0), "192A56")
        tabla_id.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tabla_id.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── SECCIÓN 1: RESUMEN EJECUTIVO ──────────────────────────────────
    _h1(doc, "1. Resumen Ejecutivo")

    nivel_riesgo = "ALTO" if any(f["nivel"] == "ALTO" for f in flags) else \
                   "MEDIO" if flags else "BAJO"
    color_nivel  = _nivel_color(nivel_riesgo)

    p_nivel = doc.add_paragraph()
    r_lbl   = p_nivel.add_run("Nivel de riesgo general: ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(11)
    r_nivel = p_nivel.add_run(f"  {nivel_riesgo}  ")
    r_nivel.bold             = True
    r_nivel.font.size        = Pt(11)
    r_nivel.font.color.rgb   = RGBColor(0xFF, 0xFF, 0xFF)
    r_nivel.font.highlight_color = None

    _kv(doc, "Total adjudicado histórico", fmt_clp(metr["total_adjudicado"]))
    _kv(doc, "Contratos registrados",      f"{metr['total_contratos']} ({metr['total_con_monto']} con monto informado)")
    _kv(doc, "Organismos compradores",     metr["organismos_unicos"])
    _kv(doc, "Monto promedio por contrato", fmt_clp(metr["monto_promedio"]))
    _kv(doc, "Alertas detectadas",         f"{len(flags)} ({sum(1 for f in flags if f['nivel']=='ALTO')} críticas)")

    # ── SECCIÓN 2: ANÁLISIS FINANCIERO ───────────────────────────────
    _h1(doc, "2. Análisis Financiero")

    _h2(doc, "Métricas principales")
    tabla_fin = doc.add_table(rows=5, cols=2)
    tabla_fin.style = "Table Grid"
    datos_fin = [
        ("Total adjudicado histórico", fmt_clp(metr["total_adjudicado"])),
        ("Monto promedio por contrato", fmt_clp(metr["monto_promedio"])),
        ("Contrato de mayor valor",    fmt_clp(metr["monto_maximo"])),
        ("Contrato de menor valor",    fmt_clp(metr["monto_minimo"])),
        ("Índice de concentración HHI", f"{metr['hhi']:.4f}  (0=diversificado / 1=monopolio)"),
    ]
    for i, (lbl, val) in enumerate(datos_fin):
        _set_cell_bg(tabla_fin.cell(i, 0), "F2F2F2")
        tabla_fin.cell(i, 0).text = lbl
        tabla_fin.cell(i, 0).paragraphs[0].runs[0].bold = True
        tabla_fin.cell(i, 1).text = val

    _h2(doc, "Evolución anual")
    por_anio = metr.get("por_anio", {})
    if por_anio:
        tabla_anio = doc.add_table(rows=len(por_anio) + 1, cols=3)
        tabla_anio.style = "Table Grid"
        headers = ["Año", "Monto adjudicado", "N° contratos"]
        for j, h in enumerate(headers):
            c = tabla_anio.cell(0, j)
            c.text = h
            _set_cell_bg(c, "192A56")
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].bold = True
        for i, (anio, v) in enumerate(sorted(por_anio.items())):
            tabla_anio.cell(i + 1, 0).text = str(anio)
            tabla_anio.cell(i + 1, 1).text = fmt_clp(v["monto"])
            tabla_anio.cell(i + 1, 2).text = str(v["contratos"])
    else:
        _body(doc, "No se encontraron datos de evolución anual.")

    _h2(doc, "Principales organismos compradores")
    top = metr.get("top_organismos", [])
    if top:
        tabla_org = doc.add_table(rows=len(top) + 1, cols=3)
        tabla_org.style = "Table Grid"
        for j, h in enumerate(["Organismo", "Monto total", "% del total"]):
            c = tabla_org.cell(0, j)
            c.text = h
            _set_cell_bg(c, "192A56")
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].bold = True
        for i, o in enumerate(top):
            tabla_org.cell(i + 1, 0).text = o["nombre"]
            tabla_org.cell(i + 1, 1).text = fmt_clp(o["monto"])
            tabla_org.cell(i + 1, 2).text = f"{o['porcentaje']}%"

    # ── SECCIÓN 3: ALERTAS ───────────────────────────────────────────
    _h1(doc, "3. Alertas y Red Flags")

    if not flags:
        _body(doc, "No se detectaron alertas de riesgo significativas en el perfil analizado.", italic=True)
    else:
        for flag in flags:
            p_flag = doc.add_paragraph()
            p_flag.paragraph_format.space_before = Pt(8)
            r_nivel = p_flag.add_run(f"[{flag['nivel']}] ")
            r_nivel.bold           = True
            r_nivel.font.color.rgb = _nivel_color(flag["nivel"])
            r_nivel.font.size      = Pt(10)
            r_titulo = p_flag.add_run(flag["titulo"])
            r_titulo.bold      = True
            r_titulo.font.size = Pt(10)

            _body(doc, flag["detalle"])
            _body(doc, f"Dato observado: {flag.get('valor', '—')}", italic=True)

    # ── SECCIÓN 4: ANÁLISIS IA ───────────────────────────────────────
    _h1(doc, "4. Análisis de Inteligencia Artificial")
    _body(doc, "El siguiente análisis fue generado por Claude Opus (Anthropic) sobre la base de los datos objetivos anteriores.", italic=True)

    # Dividir el análisis por secciones (detecta **encabezados** de markdown)
    for linea in analisis.split("\n"):
        linea = linea.strip()
        if not linea:
            continue
        if linea.startswith("**") and linea.endswith("**"):
            _h2(doc, linea.strip("*"))
        elif linea.startswith("**"):
            # Línea con negrita embebida — simplificar
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(linea.replace("**", ""))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
        else:
            _body(doc, linea)

    # ── PIE DE PÁGINA ─────────────────────────────────────────────────
    doc.add_paragraph()
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pie = p_pie.add_run(
        f"Informe generado por Radar Mercado Público · {datetime.now().strftime('%d/%m/%Y')} · "
        "Fuente de datos: API pública ChileCompra (mercadopublico.cl)"
    )
    r_pie.font.size      = Pt(8)
    r_pie.font.color.rgb = MID_GRAY
    r_pie.italic         = True

    # Serializar a bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
